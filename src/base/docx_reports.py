from docx import Document
from docx.shared import Inches
import config
import re
from src.base.logger import  Logger
import os
import pathlib

class Report_docx:

    __name = 'Класс для формирования отчётов после завершения тестирования'

    def __init__(self) -> None:
        pass

    def create_docx(self):

        tmp_file = {}
        counter_tests = 0
        counter_steps = 1

        path = str(pathlib.Path().absolute()) + '/src/base/logs/'
        log_list = os.listdir( path )
        print(log_list)
        for log_file in log_list:
            with open(path + log_file) as file:
                for line in file:
                    try:
                        if config.decorator_start_test in line:
                            counter_tests += 1
                            counter_steps = 1
                            tmp_file[counter_tests] = {}
                        elif config.indicator_test_page in line:
                            tmp_file[counter_tests][config.indicator_test_page] = str(re.search( fr'{config.indicator_test_page}.+$', line )[0]).replace(config.indicator_test_page, '') 
                        elif config.indicator_test_name in line:
                            tmp_file[counter_tests][config.indicator_test_name] = str(re.search( fr'{config.indicator_test_name}.+$', line )[0]).replace(config.indicator_test_name, '')
                        elif config.indicator_test_param in line:
                            tmp_file[counter_tests][config.indicator_test_param] = str(re.search( fr'{config.indicator_test_param}.+$', line )[0]).replace(config.indicator_test_param, '')
                        elif config.indicator_test_step in line:
                            try:
                                if type(tmp_file[counter_tests][config.indicator_test_step]) == list:
                                    tmp_file[counter_tests][config.indicator_test_step].append(str(counter_steps) + ') ' + str(re.search( fr'{config.indicator_test_step}.+$', line )[0]).replace(config.indicator_test_step, ''))
                                    counter_steps += 1
                            except KeyError:
                                tmp_file[counter_tests][config.indicator_test_step] = []
                                tmp_file[counter_tests][config.indicator_test_step].append(str(counter_steps) + ') ' + str(re.search( fr'{config.indicator_test_step}.+$', line )[0]).replace(config.indicator_test_step, ''))
                                counter_steps += 1
                        elif config.indicator_test_result_suc in line:
                            tmp_file[counter_tests][config.indicator_test_result_suc] = str(re.search( fr'{config.indicator_test_result_suc}.+$', line )[0]).replace(config.indicator_test_result_suc, '')
                            tmp_file[counter_tests][config.indicator_test_step].append(str(counter_steps) + ') ' + str(re.search( fr'{config.indicator_test_result_suc}.+$', line )[0]).replace(config.indicator_test_result_suc, ''))
                            counter_steps += 1
                        elif config.indicator_test_result_err in line:
                            tmp_file[counter_tests][config.indicator_test_result_err] = str(re.search( fr'{config.indicator_test_result_err}.+$', line )[0]).replace(config.indicator_test_result_err, '')
                            tmp_file[counter_tests][config.indicator_test_step].append(str(counter_steps) + ') ' + str(re.search( fr'{config.indicator_test_result_suc}.+$', line )[0]).replace(config.indicator_test_result_err, ''))
                            counter_steps += 1
                        elif config.indicator_test_screen in line:
                            tmp_file[counter_tests][config.indicator_test_screen] = str(re.search( fr'{config.indicator_test_screen}.+$', line )[0]).replace(config.indicator_test_screen, '')
                        elif config.decorator_end_test in line:
                            pass
                        else:
                            pass
                    except Exception as e:
                        Logger(f'Ошибка {e} при парсинге логов!').debuglog

        """ Разбиение на страницы в отчёте """
        sections = []
        for i in tmp_file:
            try:
                sections.append(tmp_file[i][config.indicator_test_page])
            except Exception as e:
                Logger(f'Ошибка {e}. Возможно какой-то тест был прерван!').debuglog
        sections = list(set(sections))
        Logger(f"Разделы попавшие в прогон {sections}").infolog

        structure_preparation = {}
        for section in sections:
            for i in tmp_file:
                if section == tmp_file[i][config.indicator_test_page]:
                    try:
                        if structure_preparation[section]:
                            for step in range(len(structure_preparation[section]['section_steps'])):
                                    
                                if structure_preparation[section]['section_steps'][step]['test_steps'] == tmp_file[i][config.indicator_test_step]:
                                    try:
                                        add_json = {
                                                'test_name' : tmp_file[i][config.indicator_test_name],
                                                'test_param' : tmp_file[i][config.indicator_test_param],
                                                'test_status' : status,
                                            }
                                        structure_preparation[section]['section_steps'][step]['cases'].append(add_json)
                                        Logger(f'Добавил 1').debuglog
                                    except Exception as e:
                                        pass

                                elif structure_preparation[section]['section_steps'][step]['test_steps'] != tmp_file[i][config.indicator_test_step]:
                                    Logger(f'Зашёл в else').debuglog
                                    error_coter = None
                                    try:
                                        tmp_file[i][config.indicator_test_result_suc]
                                        status = 'Успешно'
                                        error_coter = 1
                                    except Exception as e:
                                        pass
                                    try:
                                        tmp_file[i][config.indicator_test_result_err]
                                        status = 'Ошибка'
                                        error_coter += 1
                                    except Exception as e:
                                        pass
                                    if error_coter == None:
                                        status = 'Ошибка: Тест завершился преждевременно'

                                    added_step = []
                                    for b in structure_preparation[section]['section_steps']:
                                        added_step.append(b['test_steps'])

                                    if tmp_file[i][config.indicator_test_step] not in added_step:
                                        structure_preparation[section]['section_steps'].append(
                                            {
                                                'test_steps' : tmp_file[i][config.indicator_test_step],
                                                'cases' : [{
                                                    'test_name' : tmp_file[i][config.indicator_test_name],
                                                    'test_param' : tmp_file[i][config.indicator_test_param],
                                                    'test_status' : status,
                                                }]
                                            }
                                        )

                                        

                    except KeyError as e:
                        error_coter = None
                        try:
                            tmp_file[i][config.indicator_test_result_suc]
                            status = 'Успешно'
                            error_coter = 1
                        except Exception as e:
                            pass
                        try:
                            tmp_file[i][config.indicator_test_result_err]
                            status = 'Ошибка'
                            error_coter += 1
                        except Exception as e:
                            pass
                        if error_coter == None:
                            status = 'Ошибка: Тест завершился преждевременно'

                        structure_preparation[section] = {
                            'name' : section,
                            'section_steps' : [{
                                'test_steps' : tmp_file[i][config.indicator_test_step],
                                'cases' : [{
                                    'test_name' : tmp_file[i][config.indicator_test_name],
                                    'test_param' : tmp_file[i][config.indicator_test_param],
                                    'test_status' : status,
                                }]
                            }],
                        }

        document = Document()
        sec = document.sections[-1]
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

        pages_testing = ''
        for ss in sections:
            pages_testing += ss

        document.add_heading('Регресс MPP', 0)
        document.add_heading(f'Разделы приложения попавшие в тестирование:\n {pages_testing}', level=3)

        for section in structure_preparation:
            document.add_heading(f'Страница:\n {structure_preparation[section]["name"]}', level=5)
            
            for case in structure_preparation[section]['section_steps']:
                correct_steps = ''
                for c in case['test_steps']:
                    correct_steps += c + '\n'
                document.add_heading(f"Шаги тестирования: \n{correct_steps}", level=7)
                table = document.add_table(rows=1, cols=3, style='Table Grid')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Название'
                hdr_cells[1].text = 'Параметр'
                hdr_cells[2].text = 'Статус'
                for test in case['cases']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(test['test_name']).strip()
                    if test['test_param'] == '' or test['test_param'] == ' ':
                        row_cells[1].text = '<Пустая строка>'
                    else:
                        row_cells[1].text = str(test['test_param']).strip()
                    row_cells[2].text = str(test['test_status']).strip()
                pass

        document.save(config.name_file_for_docx_report_from_log)
